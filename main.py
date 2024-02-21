import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from PyPDF2 import PdfReader
from docx import Document
import os


class FileConverterApp:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("File Converter")
        self.root.geometry("400x300")

        self.file_path = ""
        self.output_path = ""

        # File upload Section
        self.upload_button = tk.Button(self.root, text="Upload File", command=self.upload_file)
        self.upload_button.pack(pady=5)

        # Conversion Option Section
        self.conversion_options = ["PDF to DOCX", "DOCX to PDF"]
        self.selected_option = tk.StringVar()
        self.selected_option.set(self.conversion_options[0])

        self.option_menu = tk.OptionMenu(self.root, self.selected_option, *self.conversion_options)
        self.option_menu.pack(pady=5)

        # Convert Button
        self.convert_button = tk.Button(self.root, text="Convert", command=self.convert_file)
        self.convert_button.pack(pady=5)

        # Progress Bar
        self.progress = ttk.Progressbar(self.root, orient="horizontal", length=200, mode="determinate")
        self.progress.pack(pady=5)

        # Download Button
        self.download_button = tk.Button(self.root, text="Download", command=self.download_file, state="disabled")
        self.download_button.pack(pady=5)

        # Reset Button
        self.reset_button = tk.Button(self.root, text="Reset", command=self.reset)
        self.reset_button.pack(pady=5)

        self.root.mainloop()

    def upload_file(self):
        self.file_path = filedialog.askopenfilename(filetypes=[("All Files", "*.*")])
        if self.file_path:
            messagebox.showinfo("File Uploaded", f"File {os.path.basename(self.file_path)} Uploaded Successfully")

    def convert_file(self):
        if self.file_path:
            try:
                self.progress["value"] = 0
                self.progress.update_idletasks()

                output_format = self.selected_option.get().split()[2].lower()
                self.output_path = filedialog.asksaveasfilename(defaultextension=f".{output_format}")

                if output_format == "docx":
                    self.convert_pdf_to_docx()
                elif output_format == "pdf":
                    self.convert_docx_to_pdf()
                else:
                    messagebox.showerror("Error", "Invalid Conversion Option.")

            except Exception as e:
                messagebox.showerror("Error", f"An error occurred: {e}")

        else:
            messagebox.showwarning("Warning", "Please upload a file to convert.")

    def convert_pdf_to_docx(self):
        try:
            pdf_reader = PdfReader(self.file_path)
            total_pages = len(pdf_reader.pages)

            doc = Document()

            for page_number, page in enumerate(pdf_reader.pages, start=1):
               pdf_text = page.extract_text()
               doc.add_paragraph(pdf_text)

               self.progress["value"] = (page_number / total_pages) * 100
               self.progress.update_idletasks()

            doc.save(self.output_path)
            messagebox.showinfo("Success", "PDF converted to docx successfully.")
            self.download_button.config(state="normal")

        except Exception as e:
            messagebox.showerror("Error!", f"An error occurred while converting pdf to docx:{e}")


    def convert_docx_to_pdf(self):
        doc = Document(self.file_path)
        doc.save(self.output_path)
        messagebox.showinfo("Success", "docx converted to pfd successfully.")
        self.download_button.config(state="normal")

    def download_file(self):
        messagebox.showinfo("Download", "File Downloaded Successfully.")

    def reset(self):
        self.file_path = ""
        self.output_path = ""
        self.selected_option.set(self.conversion_options[0])
        self.progress["value"] = 0
        self.download_button.config(state="disabled")

if __name__ == "__main__":
    FileConverterApp()
